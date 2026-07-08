from __future__ import annotations

import json
import sys
import time
from types import SimpleNamespace

import pytest

from src.rag import (
    build_rag_context,
    build_rag_debug,
    format_rag_sources,
    index_documents,
    query_documents,
)
from src.rag.chunker import chunk_document
from src.rag.index import build_rag_index, load_rag_index, search_rag_index
from src.rag.loader import load_document
from src.rag.schema import RagChunk, RagDocument, RagIndex, RagSearchResult
from src.rag.service import search_documents, search_documents_with_debug
from src.rag.rerank import RerankerConfig, apply_reranker, get_reranker
from src.rag.vector import cosine_similarity, embed_text, search_rag_index_hybrid, search_rag_index_vector
from src.ui.rag_panel import (
    chunk_preview_rows,
    format_rag_debug_summary,
    format_score_breakdown,
    parse_path_lines,
    sanitize_upload_name,
    summarize_rag_index,
)


def _install_fake_pypdf(
    monkeypatch,
    *,
    pages: list[str],
    encrypted: bool = False,
) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, _path: str) -> None:
            self.is_encrypted = encrypted
            self.pages = [FakePage(text) for text in pages]

    monkeypatch.setitem(sys.modules, "pypdf", SimpleNamespace(PdfReader=FakeReader))


def test_load_markdown_document_normalizes_text_and_metadata(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("  RAG notes  \r\n\r\nretrieval budget  \n", encoding="utf-8")

    document = load_document(path)

    assert document.title == "notes"
    assert document.file_type == "md"
    assert document.text == "RAG notes\n\nretrieval budget"
    assert document.metadata["size_bytes"] > 0
    assert len(document.content_hash) == 64


def test_transactional_index_save_preserves_existing_file_on_failure(monkeypatch, tmp_path):
    from src.rag import service
    from src.rag.schema import RagIndex

    target = tmp_path / "rag_index.json"
    target.write_text("old index", encoding="utf-8")

    def fail_after_temp_write(index, path):
        path.write_text("new partial index", encoding="utf-8")
        raise RuntimeError("disk interrupted")

    monkeypatch.setattr(service, "save_rag_index", fail_after_temp_write)

    with pytest.raises(RuntimeError, match="disk interrupted"):
        service._transactional_save_index(RagIndex(version=1, documents=(), chunks=()), target)

    assert target.read_text(encoding="utf-8") == "old index"
    assert not list(tmp_path.glob(".rag_index.json.*.tmp"))


def test_load_docx_document_when_python_docx_is_available(tmp_path):
    from docx import Document

    path = tmp_path / "lesson.docx"
    document = Document()
    document.add_paragraph("Session memory")
    document.add_paragraph("Citation retrieval")
    document.save(path)

    loaded = load_document(path)

    assert loaded.file_type == "docx"
    assert "Session memory" in loaded.text
    assert "Citation retrieval" in loaded.text


def test_load_pdf_document_extracts_page_text_and_metadata(tmp_path, monkeypatch):
    _install_fake_pypdf(monkeypatch, pages=["PDF retrieval text", "second page"])
    path = tmp_path / "paper.pdf"
    path.write_bytes(b"%PDF-1.4")

    document = load_document(path)

    assert document.file_type == "pdf"
    assert "[Page 1]" in document.text
    assert "PDF retrieval text" in document.text
    assert document.metadata["pdf_pages"] == 2


def test_pdf_loader_rejects_oversized_files(tmp_path):
    path = tmp_path / "large.pdf"
    path.write_bytes(b"%PDF-1.4 large")

    with pytest.raises(ValueError, match="PDF is too large"):
        load_document(path, max_pdf_bytes=4)


def test_pdf_loader_rejects_too_many_pages(tmp_path, monkeypatch):
    _install_fake_pypdf(monkeypatch, pages=["one", "two", "three"])
    path = tmp_path / "many_pages.pdf"
    path.write_bytes(b"%PDF-1.4")

    with pytest.raises(ValueError, match="PDF has too many pages"):
        load_document(path, max_pdf_pages=2)


def test_pdf_loader_rejects_encrypted_files(tmp_path, monkeypatch):
    _install_fake_pypdf(monkeypatch, pages=["secret"], encrypted=True)
    path = tmp_path / "encrypted.pdf"
    path.write_bytes(b"%PDF-1.4")

    with pytest.raises(ValueError, match="Encrypted PDF"):
        load_document(path)


def test_chunk_document_tracks_source_lines(tmp_path):
    path = tmp_path / "source.md"
    path.write_text("Intro line\nSecond line\n\nBudget line\n", encoding="utf-8")
    document = load_document(path)

    chunks = chunk_document(document, max_chars=200, overlap_chars=0)

    assert len(chunks) == 1
    assert chunks[0].start_line == 1
    assert chunks[0].end_line == 4
    assert chunks[0].metadata["char_count"] == len(chunks[0].text)


def test_build_save_load_and_query_rag_index(tmp_path):
    python_doc = tmp_path / "python.md"
    cooking_doc = tmp_path / "cooking.txt"
    python_doc.write_text("Python requests sessions reuse connections.", encoding="utf-8")
    cooking_doc.write_text("Pasta sauce needs tomatoes.", encoding="utf-8")
    index_path = tmp_path / "rag_index.json"

    index = index_documents(
        [python_doc, cooking_doc],
        index_path=index_path,
        max_chars=200,
        overlap_chars=0,
    )
    loaded = load_rag_index(index_path)
    results = query_documents("requests sessions", index_path=index_path, top_k=1)

    assert len(index.documents) == 2
    assert len(loaded.chunks) == 2
    assert results[0].chunk.source_path == str(python_doc)
    assert results[0].matched_terms == ("requests", "sessions")

    raw_index = json.loads(index_path.read_text(encoding="utf-8"))
    assert raw_index["version"] == 1


def test_chinese_query_matches_cjk_bigrams(tmp_path):
    path = tmp_path / "cn.md"
    path.write_text(
        "\u5411\u91cf\u68c0\u7d22\u5e2e\u52a9\u672c\u5730\u77e5\u8bc6\u95ee\u7b54",
        encoding="utf-8",
    )

    index = build_rag_index([path], max_chars=200, overlap_chars=0)
    results = search_rag_index(index, "\u5411\u91cf", top_k=1)

    assert results
    assert "\u5411\u91cf" in results[0].matched_terms


def test_lexical_retrieval_uses_bm25_length_normalization(tmp_path):
    concise_doc = tmp_path / "concise.md"
    long_doc = tmp_path / "repeated.md"
    concise_doc.write_text("requests sessions", encoding="utf-8")
    long_doc.write_text(("requests " * 60) + "sessions", encoding="utf-8")
    index = build_rag_index([concise_doc, long_doc], max_chars=2000, overlap_chars=0)

    results = search_rag_index(index, "requests sessions", top_k=2)

    assert results[0].chunk.source_path == str(concise_doc)
    assert results[0].score > results[1].score


def test_vector_retrieval_ranks_matching_chunk(tmp_path):
    python_doc = tmp_path / "python.md"
    cooking_doc = tmp_path / "cooking.md"
    python_doc.write_text("HTTP requests sessions reuse connections.", encoding="utf-8")
    cooking_doc.write_text("Tomato pasta sauce needs basil.", encoding="utf-8")
    index = build_rag_index([python_doc, cooking_doc], max_chars=200, overlap_chars=0)

    results = search_rag_index_vector(index, "requests connections", top_k=1)

    assert results[0].chunk.source_path == str(python_doc)
    assert 0 < results[0].score <= 1
    assert "connections" in results[0].matched_terms


def test_hybrid_retrieval_fuses_lexical_and_vector_ranks_with_rrf(tmp_path):
    python_doc = tmp_path / "python.md"
    cooking_doc = tmp_path / "cooking.md"
    python_doc.write_text("HTTP requests sessions reuse connections.", encoding="utf-8")
    cooking_doc.write_text("Tomato pasta sauce needs basil.", encoding="utf-8")
    index = build_rag_index([python_doc, cooking_doc], max_chars=200, overlap_chars=0)

    results = search_rag_index_hybrid(index, "requests connections", top_k=1)

    assert results[0].chunk.source_path == str(python_doc)
    assert 0 < results[0].score < 0.05
    assert "requests" in results[0].matched_terms


def test_query_documents_supports_retrieval_modes(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("Local vector retrieval keeps cited chunks.", encoding="utf-8")
    index_path = tmp_path / "rag_index.json"
    index_documents([path], index_path=index_path, max_chars=200, overlap_chars=0)

    lexical_results = query_documents("cited chunks", index_path=index_path, retrieval_mode="lexical")
    hybrid_results = query_documents("cited chunks", index_path=index_path, retrieval_mode="hybrid")
    vector_results = query_documents("cited chunks", index_path=index_path, retrieval_mode="vector")
    backend_results = query_documents(
        "cited chunks",
        index_path=index_path,
        retrieval_mode="backend_vector",
    )

    assert lexical_results
    assert hybrid_results
    assert vector_results
    assert backend_results


def test_query_documents_rejects_unknown_retrieval_mode(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("Local retrieval.", encoding="utf-8")
    index_path = tmp_path / "rag_index.json"
    index_documents([path], index_path=index_path, max_chars=200, overlap_chars=0)

    with pytest.raises(ValueError, match="Unsupported RAG retrieval mode"):
        query_documents("retrieval", index_path=index_path, retrieval_mode="semantic")


def test_search_documents_filters_duplicates_metadata_and_source_diversity():
    index = RagIndex(
        version=1,
        documents=(
            RagDocument(
                source_path="python.md",
                title="python",
                text="alpha retrieval session",
                content_hash="doc-a",
                file_type="md",
                document_id="doc-a",
                metadata={"course": "python"},
            ),
            RagDocument(
                source_path="java.md",
                title="java",
                text="alpha retrieval session",
                content_hash="doc-b",
                file_type="md",
                document_id="doc-b",
                metadata={"course": "java"},
            ),
        ),
        chunks=(
            RagChunk(
                chunk_id="a1",
                document_hash="doc-a",
                source_path="python.md",
                title="python",
                text="alpha retrieval session",
                chunk_index=0,
                start_line=1,
                end_line=1,
                document_id="doc-a",
            ),
            RagChunk(
                chunk_id="a2",
                document_hash="doc-a",
                source_path="python.md",
                title="python",
                text="alpha retrieval session",
                chunk_index=1,
                start_line=2,
                end_line=2,
                document_id="doc-a",
            ),
            RagChunk(
                chunk_id="b1",
                document_hash="doc-b",
                source_path="java.md",
                title="java",
                text="alpha retrieval stream",
                chunk_index=2,
                start_line=1,
                end_line=1,
                document_id="doc-b",
            ),
        ),
    )

    filtered = search_documents(
        index,
        "alpha retrieval",
        retrieval_mode="lexical",
        top_k=3,
        metadata_filters={"course": "python"},
    )
    diverse = search_documents(
        index,
        "alpha retrieval",
        retrieval_mode="lexical",
        top_k=3,
        max_chunks_per_source=1,
    )

    assert [result.chunk.chunk_id for result in filtered] == ["a1"]
    assert [result.chunk.document_id for result in diverse] == ["doc-a", "doc-b"]


def test_search_documents_with_debug_records_backend_vector_latency(monkeypatch, tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("Backend vector retrieval keeps diagnostics.", encoding="utf-8")
    index = build_rag_index([path], max_chars=200, overlap_chars=0)

    class FakeBackend:
        name = "fake"

        def query(self, index, query, *, top_k=5, min_score=0.05):
            time.sleep(0.002)
            return [
                RagSearchResult(
                    chunk=index.chunks[0],
                    score=0.9,
                    matched_terms=("backend",),
                )
            ]

    from src.rag import service

    monkeypatch.setattr(service, "get_vector_backend_from_env", lambda: FakeBackend())

    diagnostics = search_documents_with_debug(
        index,
        "backend diagnostics",
        retrieval_mode="backend_vector",
        top_k=1,
    )

    stage = diagnostics.debug["stages"][0]
    assert stage["name"] == "backend_vector"
    assert stage["elapsed_ms"] > 0
    assert diagnostics.debug["post_filter"]["output_count"] == 1


def test_lexical_overlap_reranker_reorders_candidates_by_query_terms():
    weak = RagSearchResult(
        chunk=RagChunk(
            chunk_id="weak",
            document_hash="doc-a",
            source_path="weak.md",
            title="weak",
            text="alpha only",
            chunk_index=0,
            start_line=1,
            end_line=1,
        ),
        score=0.9,
        matched_terms=("alpha",),
    )
    strong = RagSearchResult(
        chunk=RagChunk(
            chunk_id="strong",
            document_hash="doc-b",
            source_path="strong.md",
            title="strong",
            text="alpha beta gamma",
            chunk_index=1,
            start_line=1,
            end_line=1,
        ),
        score=0.5,
        matched_terms=("alpha", "beta", "gamma"),
    )

    outcome = apply_reranker(
        "alpha beta gamma",
        [weak, strong],
        config=RerankerConfig(name="lexical_overlap", top_n=2, latency_budget_ms=250),
    )

    assert get_reranker("lexical_overlap").name == "lexical_overlap"
    assert [result.chunk.chunk_id for result in outcome.results] == ["strong", "weak"]
    assert outcome.stage["name"] == "reranker:lexical_overlap"
    assert outcome.stage["within_latency_budget"] is True
    assert outcome.stage["within_cost_budget"] is True


def test_search_documents_with_debug_includes_reranker_stage(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("alpha beta gamma\n\nalpha only", encoding="utf-8")
    index = build_rag_index([path], max_chars=18, overlap_chars=0)

    diagnostics = search_documents_with_debug(
        index,
        "alpha beta gamma",
        retrieval_mode="hybrid",
        top_k=2,
        reranker="lexical_overlap",
        rerank_top_n=2,
    )

    assert any(stage["name"] == "reranker:lexical_overlap" for stage in diagnostics.debug["stages"])


def test_build_rag_debug_explains_hybrid_scores(tmp_path):
    python_doc = tmp_path / "python.md"
    cooking_doc = tmp_path / "cooking.md"
    python_doc.write_text("HTTP requests sessions reuse connections.", encoding="utf-8")
    cooking_doc.write_text("Tomato pasta sauce needs basil.", encoding="utf-8")
    index = build_rag_index([python_doc, cooking_doc], max_chars=200, overlap_chars=0)
    results = search_rag_index_hybrid(index, "requests connections", top_k=1)

    debug = build_rag_debug(
        index,
        "requests connections",
        results,
        retrieval_mode="hybrid",
        top_k=1,
        min_score=0.01,
    )

    assert debug["candidate_count"] == 2
    assert debug["returned_count"] == 1
    assert [stage["name"] for stage in debug["stages"]] == ["lexical_bm25", "local_vector"]
    assert all(stage["candidate_count"] == 2 for stage in debug["stages"])
    assert all(stage["elapsed_ms"] >= 0 for stage in debug["stages"])
    assert debug["query_terms"] == ["connections", "requests"]
    breakdown = debug["results"][0]["score_breakdown"]
    assert breakdown["fusion"] == "rrf"
    assert breakdown["rrf_k"] == 60
    assert breakdown["lexical_rank"] == 1
    assert breakdown["lexical_rrf"] > 0
    assert breakdown["vector_rank"] >= 1
    assert breakdown["vector_rrf"] > 0
    assert breakdown["combined_score"] == results[0].score
    assert breakdown["lexical_score"] > 0
    assert breakdown["vector_score"] > 0


def test_rag_panel_index_summary_and_chunk_preview(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("First retrieval paragraph.\n\nSecond retrieval paragraph.", encoding="utf-8")
    index = build_rag_index([path], max_chars=200, overlap_chars=0)

    summary = summarize_rag_index(index)
    preview_rows = chunk_preview_rows(index)

    assert summary["documents"] == 1
    assert summary["chunks"] == 1
    assert summary["document_rows"][0]["title"] == "notes"
    assert summary["document_rows"][0]["chunk_count"] == 1
    assert len(summary["document_rows"][0]["content_hash"]) == 8
    assert preview_rows[0]["lines"] == "L1-L3"
    assert "retrieval paragraph" in preview_rows[0]["preview"]


def test_rag_panel_formats_debug_summary_and_breakdown():
    debug = {
        "retrieval_mode": "hybrid",
        "top_k": 3,
        "min_score": 0.01,
        "candidate_count": 8,
        "returned_count": 2,
        "query_terms": ["rag", "debug"],
    }
    result_debug = {
        "score_breakdown": {
            "fusion": "rrf",
            "rrf_k": 60,
            "lexical_rank": 1,
            "lexical_rrf": 0.016393,
            "lexical_score": 3.5,
            "lexical_normalized": 1.0,
            "vector_rank": 2,
            "vector_rrf": 0.016129,
            "vector_score": 0.25,
            "combined_score": 0.032522,
        }
    }

    assert format_rag_debug_summary(debug) == (
        "mode=hybrid; top_k=3; min_score=0.01; candidates=8; returned=2; terms=debug, rag"
    )
    assert "fusion=rrf" in format_score_breakdown(result_debug)
    assert "combined_score=0.033" in format_score_breakdown(result_debug)


def test_build_rag_debug_marks_empty_queries(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("Local retrieval.", encoding="utf-8")
    index = build_rag_index([path], max_chars=200, overlap_chars=0)

    debug = build_rag_debug(
        index,
        "",
        [],
        retrieval_mode="lexical",
        top_k=3,
        min_score=0.01,
    )

    assert debug["empty_query"] is True
    assert debug["candidate_count"] == 1
    assert debug["returned_count"] == 0
    assert debug["stages"][0]["name"] == "lexical_bm25"


def test_local_hash_embeddings_are_deterministic():
    left = embed_text("requests session reuse")
    right = embed_text("requests session reuse")

    assert left == right
    assert cosine_similarity(left, right) == pytest.approx(1.0)


def test_build_rag_context_includes_citations_and_limits_text(tmp_path):
    path = tmp_path / "architecture.md"
    path.write_text("RAG architecture uses cited chunks for answers.", encoding="utf-8")
    index = build_rag_index([path], max_chars=200, overlap_chars=0)
    results = search_rag_index(index, "cited chunks", top_k=1)

    context = build_rag_context(results, max_chars=500)

    assert "[1] architecture" in context
    assert f"{path}:L1-L1" in context
    assert "cited chunks" in context
    assert len(context) <= 500


def test_format_rag_sources_lists_file_lines_and_terms(tmp_path):
    path = tmp_path / "architecture.md"
    path.write_text("RAG architecture uses cited chunks for answers.", encoding="utf-8")
    index = build_rag_index([path], max_chars=200, overlap_chars=0)
    results = search_rag_index(index, "cited chunks", top_k=1)

    source_block = format_rag_sources(results)

    assert "[1] architecture" in source_block
    assert f"{path}:L1-L1" in source_block
    assert "matched=chunks, cited" in source_block


def test_empty_rag_context_is_explicit():
    assert build_rag_context([]) == "No relevant local documents retrieved."


def test_empty_rag_sources_are_blank():
    assert format_rag_sources([]) == ""


def test_rag_upload_name_is_sanitized():
    assert sanitize_upload_name("../unsafe lesson.md") == "unsafe_lesson.md"
    assert sanitize_upload_name("资料 01.docx") == "01.docx"
    assert sanitize_upload_name("...") == "document"


def test_parse_path_lines_ignores_blank_lines():
    paths = parse_path_lines(' "notes.md" \n\n C:/tmp/lesson.txt ')

    normalized = [str(path).replace("\\", "/") for path in paths]
    assert normalized == ["notes.md", "C:/tmp/lesson.txt"]
