from __future__ import annotations

import hashlib
from collections.abc import Sequence
from pathlib import Path

from src.rag.schema import RagDocument

SUPPORTED_TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | {".docx", ".pdf"}
DEFAULT_MAX_PDF_BYTES = 10 * 1024 * 1024
DEFAULT_MAX_PDF_PAGES = 30
DEFAULT_MAX_PDF_CHARS = 120_000


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _normalize_text(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in normalized.split("\n")]
    paragraphs: list[str] = []
    blank_seen = False
    for line in lines:
        if line.strip():
            paragraphs.append(line)
            blank_seen = False
        elif paragraphs and not blank_seen:
            paragraphs.append("")
            blank_seen = True
    return "\n".join(paragraphs).strip()


def _read_text_path(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_docx_path(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required to load .docx files") from exc

    document = Document(str(path))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs]
    return "\n\n".join(part for part in parts if part)


def _read_pdf_path(
    path: Path,
    *,
    max_bytes: int = DEFAULT_MAX_PDF_BYTES,
    max_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_chars: int = DEFAULT_MAX_PDF_CHARS,
) -> tuple[str, dict[str, int]]:
    if max_bytes <= 0:
        raise ValueError("max_bytes must be positive")
    if max_pages <= 0:
        raise ValueError("max_pages must be positive")
    if max_chars <= 0:
        raise ValueError("max_chars must be positive")
    if path.stat().st_size > max_bytes:
        raise ValueError(f"PDF is too large: {path}")

    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("pypdf is required to load .pdf files") from exc

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError(f"Encrypted PDF is not supported: {path}")
    page_count = len(reader.pages)
    if page_count > max_pages:
        raise ValueError(f"PDF has too many pages: {page_count} > {max_pages}")

    parts: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"[Page {page_index}]\n{text.strip()}")
        if sum(len(part) for part in parts) > max_chars:
            raise ValueError(f"PDF extracted text is too long: {path}")

    return "\n\n".join(parts), {"pdf_pages": page_count}


def load_document(
    path: str | Path,
    *,
    max_pdf_bytes: int = DEFAULT_MAX_PDF_BYTES,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_pdf_chars: int = DEFAULT_MAX_PDF_CHARS,
) -> RagDocument:
    """Load one local document into a normalized RAG document."""
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(str(source))

    suffix = source.suffix.lower()
    loader_metadata: dict[str, int] = {}
    if suffix in SUPPORTED_TEXT_EXTENSIONS:
        raw_text = _read_text_path(source)
    elif suffix == ".docx":
        raw_text = _read_docx_path(source)
    elif suffix == ".pdf":
        raw_text, loader_metadata = _read_pdf_path(
            source,
            max_bytes=max_pdf_bytes,
            max_pages=max_pdf_pages,
            max_chars=max_pdf_chars,
        )
    else:
        raise ValueError(f"Unsupported RAG document type: {suffix or '<none>'}")

    text = _normalize_text(raw_text)
    if not text:
        raise ValueError(f"RAG document is empty: {source}")

    stat = source.stat()
    return RagDocument(
        source_path=str(source),
        title=source.stem,
        text=text,
        content_hash=_sha256_text(text),
        file_type=suffix.lstrip("."),
        metadata={
            "size_bytes": stat.st_size,
            "mtime_ns": stat.st_mtime_ns,
            **loader_metadata,
        },
    )


def load_documents(paths: Sequence[str | Path]) -> list[RagDocument]:
    return [load_document(path) for path in paths]
