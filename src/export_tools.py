"""数据导出 — 学习报告、群聊记录、项目状态"""

from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
EXPORT_DIR = ROOT / "exports"


def _ensure_export_dir():
    EXPORT_DIR.mkdir(exist_ok=True)


def export_session_report(messages: list[dict], meta: dict | None = None) -> Path:
    """导出本轮学习报告为 Markdown"""
    _ensure_export_dir()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = EXPORT_DIR / f"session_report_{ts}.md"
    lines = [f"# 学习报告 {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"]
    if meta:
        lines.append(f"- 角色: {meta.get('role', '-')}")
        lines.append(f"- 模式: {meta.get('mode', '-')}")
        lines.append(f"- 模型: {meta.get('model', '-')}\n")
    for m in messages:
        role = "用户" if m["role"] == "user" else "Agent"
        lines.append(f"## {role}\n{m['content']}\n")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def export_session_docx(messages: list[dict], meta: dict | None = None) -> Path:
    """导出本轮学习报告为 docx"""
    _ensure_export_dir()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = EXPORT_DIR / f"session_report_{ts}.docx"
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("学习报告", 0)
        doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
        if meta:
            doc.add_paragraph(
                f"角色: {meta.get('role', '-')}  模式: {meta.get('mode', '-')}  模型: {meta.get('model', '-')}"
            )
        for m in messages:
            role = "用户" if m["role"] == "user" else "Agent"
            doc.add_heading(role, 2)
            doc.add_paragraph(m["content"])
        doc.save(str(path))
    except ImportError:
        # fallback to markdown if python-docx missing
        return export_session_report(messages, meta)
    return path


def export_wechat_records() -> Path:
    """导出微信群完整记录"""
    _ensure_export_dir()
    group = ROOT / "chat" / "wechat_group.md"
    if not group.is_file():
        raise FileNotFoundError("wechat_group.md 不存在")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = EXPORT_DIR / f"wechat_records_{ts}.md"
    content = group.read_text(encoding="utf-8")
    path.write_text(content, encoding="utf-8")
    return path


def export_project_status() -> Path:
    """导出当前项目状态"""
    _ensure_export_dir()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = EXPORT_DIR / f"project_status_{ts}.md"
    from src.memory import read_memory_bundle

    bundle = read_memory_bundle()
    lines = [f"# 项目状态导出 {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"]
    for name in ["summary.md", "progress.md", "current_focus.md", "task_board.md"]:
        content = bundle.get(name, "")
        if content and not content.startswith("[文件不存在"):
            lines.append(f"## {name}\n{content}\n")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def export_session_archive() -> Path:
    """导出 session_archive"""
    _ensure_export_dir()
    archive = ROOT / "logs" / "session_archive.md"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = EXPORT_DIR / f"session_archive_{ts}.md"
    content = archive.read_text(encoding="utf-8") if archive.is_file() else "暂无归档"
    path.write_text(content, encoding="utf-8")
    return path
